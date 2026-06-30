import pdfplumber
import docx
import re


def extract_text_from_pdf(file_path):
    text = ""

    with pdfplumber.open(file_path) as pdf:
        
        for page in pdf.pages:
            extracted = page.extract_text()

            if extracted :
                text += extracted + "\n"

    return text

def extract_text_from_docx(file_path):
    document = docx.Document(file_path)
    text = ""

    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"

    return text

def clean_resume_text(text):
    text = re.sub(r"\s+", " ", text)

    text = text.strip()

    return text